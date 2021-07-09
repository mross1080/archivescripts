from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
doc = Document("../archive_data/translationsen.docx")

testimonies = {
    "Living on an island": {"question": "Would you go back to living on an island without a border and why?",
                            "id" : "living_on_an_island",
                            "data": []},
    "Columbus": {"question": "What did you learn about Columbus in school that you consider to be true history?",
                  "id" : "columbus",
                 "data": []},
    "Our Lady of Mercedes": {"question": "Why do you think that the Virgin Mercedes protected the Spaniards and not the Indigenous people in the battle of El Santo Cerro in 1495?",
                             "id": "our_lady_of_mercedes",
                             "data": []},
    "Enslaved People": {"question": "What do you know about enslaved Black people in the Dominican Republic?",
                        "id": "enslaved_people",
                        "data": []},
    "Plantains": {"question": "Do you eat plantains? Did you know that mangú is a food popularized by enslaved Black people in the Dominican Republic?",
                  "id": "plantains",
                  "data": []},

    "Salsa, Merengue, Bachata": {"question": "When and where do you dance to music with African influences; Salsa, Merengue, Bachata?",
                                 "id": "salsa_merengue_bachata",
                                 "data": []},
    "ID Card": {"question": "On your old identification, were you referred to as an Indio? Why?",
                "id": "id_card",
                "data": []},
    "V Century": {"question": "Why do you think the celebration of the V Centenary was important for the government and businessmen of the Dominican Republic?",
                  "id": "v_century",
                  "data": []},
    "Peña Gómez": {"question": "Do you think Peña Gómez had the right to be president, and why?",
                   "id": "pena_gomez",
                   "data": []},

}

# questions = [q["question"] for q in testimonies]

formatted_questions = []
audio_upcoming = False
def format_cell_to_archive_dict(cell_text):
    archive_for_testimony = []
    audio_upcoming = False
    msg_body = ""
    msg_sender = ""
    msg_file_type = "text"
    for line in cell_text.split("\n"):

        if (line != ""):
            # print("Raw line : " ,line)

            if (audio_upcoming):
                msg_body = line
                audio_upcoming = False

            # if "Audio" in line:
            #     print(line.split(":"))

            if (len(line.split(":")) >= 2) :
                # print(line)
                split = line.split(":")
                msg_body = split[1]

                msg_sender = "participant"

                if "L: " in line:
                    msg_sender = "admin"

                # print("{} : {}".format(msg_sender,msg_body))

                if "Audio" in line:
                    audio_upcoming = True
                    # print("next line is audio")


            if not audio_upcoming and msg_body != "":
                message = {"msg_file_type":msg_file_type,"msg_body":msg_body,"msg_sender": msg_sender}
                archive_for_testimony.append(message)
                # print(message)
    return archive_for_testimony

def extract_english_translations():
    print("STARTING ENGLISH TRANSLATIONS")
    current_category = "Living on an island"
    final_archive = []
    for question in testimonies:
        formatted_questions.append(question)
    
    doc = Document("../archive_data/translationsen.docx")
    # print(formatted_questions)
    for row in doc.tables[0].rows:
        text = row.cells[1].text.strip()
        # print(text)
        # print("QUES" )
        if "QUES" in text:
            # print("COLOM")
            # We have the named question so start with the new entry
            print(text)
        if "CATEGORY" in text:
            print("Parsed {} Testimonies For Category {}".format(len(testimonies[current_category]["data"]),current_category))
            current_category = text.split("CATEGORY: ")[1]
            print("Currently aggregating testimonies for ", current_category)
            testimonies[current_category]["name"] = current_category
        if "Testimony " in text:
            # print(text.split("\n")[0])

            testimonies[current_category]["data"].append(format_cell_to_archive_dict(text))

        # print(text.split("\n"))
    print("Finished Aggregations ")
    # print(testimonies)
    for archive in testimonies:
        final_archive.append(testimonies[archive])
    # print(final_archive)
    return final_archive




# extract_english_translations()
